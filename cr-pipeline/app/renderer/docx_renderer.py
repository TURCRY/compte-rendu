from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_paragraph(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def render_docx(j: dict) -> BytesIO:
    doc = Document()
    add_heading(doc, "Compte rendu de réunion", level=0)
    add_paragraph(doc, f"Date    {j.get('date') or '—'}")
    add_paragraph(doc, f"Link    {j.get('link') or '—'}\n")

    add_heading(doc, "Résumé", level=1)
    add_paragraph(doc, (j.get("resume") or "").strip())

    add_heading(doc, "Ordre du jour", level=1)
    for item in j.get("ordre_du_jour") or []:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Thèmes abordés", level=1)
    for t in j.get("themes_abordes") or []:
        add_heading(doc, t.get("titre","(Sans titre)"), level=2)
        for p in t.get("synthese") or []:
            doc.add_paragraph(p, style="List Bullet")
        idx = t.get("indices_source") or []
        if idx:
            add_paragraph(doc, "Indices de source", bold=True)
            for ix in idx:
                tc = ix.get("timecode","")
                sp = ix.get("speaker","")
                ex = ix.get("extrait","")
                doc.add_paragraph(f"{tc} — {sp} : {ex}", style="List Bullet")

    add_heading(doc, "Actions", level=1)
    actions = j.get("actions") or []
    if actions:
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Action"; hdr[1].text = "Responsable"; hdr[2].text = "Échéance"; hdr[3].text = "Commentaire"
        for a in actions:
            row = table.add_row().cells
            row[0].text = a.get("action","")
            row[1].text = a.get("responsable","")
            row[2].text = (a.get("echeance") or "—")
            row[3].text = (a.get("commentaire") or "—")
    else:
        add_paragraph(doc, "—")

    add_heading(doc, "Perspectives", level=1)
    for p in j.get("perspectives") or []:
        add_paragraph(doc, "Problème :", bold=True)
        add_paragraph(doc, p.get("probleme",""))
        add_paragraph(doc, "Solution envisagée / décidée :", bold=True)
        add_paragraph(doc, p.get("solution",""))

    bio = BytesIO()
    doc.save(bio); bio.seek(0)
    return bio
