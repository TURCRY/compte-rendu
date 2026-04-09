from io import BytesIO
from docx import Document
from datetime import datetime, date
import os
import re


FR_MONTHS = [
    "", "janvier", "février", "mars", "avril", "mai", "juin",
    "juillet", "août", "septembre", "octobre", "novembre", "décembre"
]

def format_date_fr(value) -> str:
    if not value:
        return ""

    # 1) Si c'est déjà un objet date/datetime
    if isinstance(value, datetime):
        dt = value.date()
        return f"{dt.day} {FR_MONTHS[dt.month]} {dt.year}"
    if isinstance(value, date):
        return f"{value.day} {FR_MONTHS[value.month]} {value.year}"

    # 2) Sinon, on traite comme une chaîne
    s = str(value).strip()

    # Accepter ISO date+heure (ex: 2025-07-03T00:00:00)
    if "T" in s:
        s = s.split("T", 1)[0]
    # Accepter "2025-07-03 00:00:00"
    if " " in s:
        s = s.split(" ", 1)[0]

    try:
        dt = datetime.strptime(s, "%Y-%m-%d").date()
        return f"{dt.day} {FR_MONTHS[dt.month]} {dt.year}"
    except Exception:
        return str(value)


def validate_payload(data: dict) -> tuple[bool, str | None]:
    required_top = ["resume", "themes_abordes", "actions", "perspectives"]
    for key in required_top:
        if key not in data:
            return False, f"Champ manquant dans le JSON : {key}"
    return True, None

def collect_global_demands(data):
    demands = []

    for k in ("demandes_documents_globales", "demandes_documents"):
        v = data.get(k) or []
        if isinstance(v, list):
            demands.extend(v)

    for sujet in (data.get("sujets") or []):
        if not isinstance(sujet, dict):
            continue
        for k in ("demandes_documents", "demandes_documents_hors_sujet"):
            v = sujet.get(k) or []
            if isinstance(v, list):
                for d in v:
                    if d is None:
                        continue
                    if isinstance(d, dict):
                        d2 = dict(d)
                    elif isinstance(d, str):
                        # tolérance : demande sous forme texte
                        d2 = {"objet": d}
                    else:
                        # type inattendu => on ignore
                        continue

                    d2["_sujet"] = sujet.get("numero")
                    demands.append(d2)

    return demands


def add_text_as_paragraphs(doc, text: str, style: str | None = None):
    t = (text or "").replace("\r\n", "\n").strip()
    if not t:
        return

    # 1) Paragraphes explicites
    if "\n\n" in t:
        parts = [p.strip() for p in t.split("\n\n") if p.strip()]
    elif "\n" in t:
        parts = [p.strip() for p in t.split("\n") if p.strip()]
    else:
        # 2) Découpage par phrases (prudent)
        parts = [p.strip() for p in re.split(r"(?<=[.!?;:])\s+", t) if p.strip()]

    for p in parts:
        try:
            doc.add_paragraph(p, style=style) if style else doc.add_paragraph(p)
        except Exception:
            doc.add_paragraph(p)


def render_markdown(data: dict) -> str:
    md = []
    md.append("# Compte rendu\n")
    md.append(f"## Résumé\n{data.get('resume','')}\n")

    md.append("## Thèmes abordés\n")
    for th in (data.get("themes_abordes") or []):
        if not isinstance(th, dict):
            continue
        md.append(f"### {th.get('titre','')}")
        for p in (th.get("synthese") or []):
            if p is None:
                continue
            md.append(f"- {str(p)}")

    if data.get("actions"):
        md.append("\n## Actions\n")
        for a in (data.get("actions") or []):
            if not isinstance(a, dict):
                continue            
            line = f"- **{a.get('action','')}**, resp. : {a.get('responsable','')}"
            if a.get("echeance"):
                line += f", échéance : {a['echeance']}"
            md.append(line)

    if data.get("perspectives"):
        md.append("\n## Perspectives\n")
        for p in (data.get("perspectives") or []):
            if not isinstance(p, dict):
                continue           
            md.append(f"- Problème : {p.get('probleme','')}")
            md.append(f"  Solution : {p.get('solution','')}")

    return "\n".join(md)

def render_docx(data: dict) -> BytesIO:
    template = os.getenv("DOCX_TEMPLATE")
    doc = Document(template) if template and os.path.exists(template) else Document()

    # -------------------------------------------------
    # Titre + informations générales
    # -------------------------------------------------
    doc.add_heading("Compte rendu d’expertise", level=1)

    doc.add_heading("Informations générales", level=2)
    doc.add_paragraph(f"Date : {format_date_fr(data.get('date'))}")
    doc.add_paragraph(f"Lien : {data.get('link') or ''}")

    # -------------------------------------------------
    # Résumé
    # -------------------------------------------------
    doc.add_heading("Résumé", level=2)
    resume = (data.get("resume") or "").replace("\r\n", "\n").strip()

    blocs = []
    if resume:
        blocs = [b.strip() for b in resume.split("\n\n") if b.strip()]
        if len(blocs) <= 1 and "\n" in resume:
            blocs = [b.strip() for b in resume.split("\n") if b.strip()]

    # Déduplication (évite répétitions exactes)
    dedup = []
    prev = None
    for b in blocs:
        if b != prev:
            dedup.append(b)
        prev = b
    blocs = dedup


    style_names = {s.name for s in doc.styles}
    use_resume_style = "Resume" in style_names

    for bloc in blocs:
        p = doc.add_paragraph(bloc)
        if use_resume_style:
            p.style = "Resume"


    # -------------------------------------------------
    # Ordre du jour
    # -------------------------------------------------
    doc.add_heading("Ordre du jour", level=2)
    for item in (data.get("ordre_du_jour") or []):
        if item is None:
            continue
        try:
            doc.add_paragraph(str(item), style="Enumération")
        except Exception:
            doc.add_paragraph(str(item))

    # -------------------------------------------------
    # Thèmes abordés
    # -------------------------------------------------
    doc.add_heading("Thèmes abordés", level=2)
    for th in (data.get("themes_abordes") or []):
        if not isinstance(th, dict):
            continue        
        doc.add_heading(th.get("titre", ""), level=3)

        for s in (th.get("synthese") or []):
            if s is None:
                continue
            try:
                doc.add_paragraph(str(s), style="Enumération")
            except Exception:
                doc.add_paragraph(str(s))

        idxs = th.get("indices_source") or []
        if isinstance(idxs, dict):
            idxs = [idxs]
        elif isinstance(idxs, str):
            idxs = [idxs]
        elif not isinstance(idxs, list):
            idxs = []

        def _is_valid_timecode(value) -> bool:
            if value in (None, "", "None"):
                return False
            s = str(value).strip()
            return bool(re.match(r"^\d{2}:\d{2}:\d{2}$", s))

        def _clean_text(value) -> str:
            if value in (None, "", "None"):
                return ""
            return str(value).strip()

        for idx in idxs:
            if idx is None:
                continue

            txt = ""

            if isinstance(idx, str):
                s = idx.strip()
                # On ignore les chaînes manifestement bruitées
                if not s or "None" in s:
                    continue
                txt = s

            elif isinstance(idx, dict):
                timecode = _clean_text(idx.get("timecode"))
                speaker = _clean_text(idx.get("speaker"))
                extrait = _clean_text(idx.get("extrait"))

                # Pas d’extrait utile => on n’affiche rien
                if not extrait:
                    continue

                # Si le timecode est invalide, on n’affiche pas la source
                # et on garde seulement l’extrait textuel
                if not _is_valid_timecode(timecode):
                    txt = extrait
                elif speaker:
                    txt = f"({timecode} – {speaker}) {extrait}"
                else:
                    txt = f"({timecode}) {extrait}"

            else:
                continue

            if not txt:
                continue

            try:
                doc.add_paragraph(txt, style="Citation intense")
            except Exception:
                doc.add_paragraph(txt)


    # -------------------------------------------------
    # Analyse par sujet
    # -------------------------------------------------
    sujets = data.get("sujets") or []
    if sujets:
        doc.add_heading("Analyse par sujet", level=2)

        for sujet in sujets:
            if not isinstance(sujet, dict):
                continue

            doc.add_heading(
                f"Sujet n°{sujet.get('numero','')} – {sujet.get('titre','')}",
                level=3
            )

            if sujet.get("localisation"):
                doc.add_paragraph(f"Localisation : {sujet.get('localisation')}")
            if sujet.get("description"):
                doc.add_paragraph(f"Description : {sujet.get('description')}")


            # Avis / Observations
            avis = sujet.get("avis_participants")

            # --- Normalisation STRUCTURELLE ---
            if avis is None:
                avis = []
            elif isinstance(avis, dict):
                avis = [avis]
            elif isinstance(avis, str):
                avis = [{"nom": "", "role": "", "resume": avis}]
            elif not isinstance(avis, list):
                avis = []

            avis_sans_expert, avis_expert = [], []

            for av in avis:
                # --- Normalisation ÉLÉMENT ---
                if av is None:
                    continue
                if isinstance(av, str):
                    av = {"nom": "", "role": "", "resume": av}
                if not isinstance(av, dict):
                    continue
                
                # Compat anciens schémas : avis/commentaire/texte -> resume
                if "resume" not in av or av.get("resume") is None:
                    for k in ("avis", "commentaire", "texte"):
                        v = av.get(k)
                        if isinstance(v, str) and v.strip():
                            av["resume"] = v
                            break
                    av.setdefault("resume", "")

                av.setdefault("nom", "")
                av.setdefault("role", "")


                role = (av.get("role") or "").lower()
                nom  = (av.get("nom") or "").lower()

                if "expert" in role or "expert" in nom:
                    avis_expert.append(av)
                else:
                    avis_sans_expert.append(av)

            # --- Rendu ---
            if avis_sans_expert:
                doc.add_heading("Avis des participants", level=4)
                for av in avis_sans_expert:
                    line = f"{av.get('nom','')} ({av.get('role','')}) : {av.get('resume','')}"
                    try:
                        doc.add_paragraph(line, style="Enumération")
                    except Exception:
                        doc.add_paragraph(line)

            if avis_expert:
                doc.add_heading("Observations de l’expert", level=4)
                for av in avis_expert:
                    txt = av.get("resume", "")
                    if txt is not None and not isinstance(txt, str):
                        txt = str(txt)
                    add_text_as_paragraphs(doc, txt, style="Normal")


            se = sujet.get("synthese_echanges")
            if se is not None and not isinstance(se, str):
                se = str(se)

            ce = sujet.get("conclusion_expert")
            if ce is not None and not isinstance(ce, str):
                ce = str(ce)

            if se:
                doc.add_heading("Synthèse des échanges", level=4)
                add_text_as_paragraphs(doc, se, style="Normal")

            if ce:
                doc.add_heading("Conclusion de l’expert", level=4)
                add_text_as_paragraphs(doc, ce, style="Normal")

            # Demandes par sujet
            dds = []
            for key in ("demandes_documents", "demandes_documents_hors_sujet"):
                v = sujet.get(key)
                if isinstance(v, list):
                    dds.extend(v)
                elif isinstance(v, dict):
                    dds.append(v)
                elif isinstance(v, str) and v.strip():
                    dds.append({"objet": v})

            if dds:
                doc.add_heading("Demandes de documents", level=4)

                for d in dds:
                    if d is None:
                        continue
                    if isinstance(d, str):
                        d = {"objet": d}
                    if not isinstance(d, dict):
                        continue

                    objet = (d.get("objet") or "").strip()
                    if not objet:
                        continue

                    parts = [objet]
                    for label, key in (
                        ("demandeur", "demandeur"),
                        ("destinataire", "destinataire"),
                        ("échéance", "echeance"),
                    ):
                        val = d.get(key)
                        if val not in (None, "", "None"):
                            parts.append(f"{label} : {val}")

                    line = parts[0] if len(parts) == 1 else f"{parts[0]} ({', '.join(parts[1:])})"
                    try:
                        doc.add_paragraph(line, style="Enumération")
                    except Exception:
                        doc.add_paragraph(line)
            

    # -------------------------------------------------
    # Actions
    # -------------------------------------------------
    doc.add_heading("Actions", level=2)
    for a in (data.get("actions") or []):
        if not isinstance(a, dict):
            continue
        action = (a.get("action") or "").strip()
        if not action:
            continue

        parts = [action]
        if a.get("responsable") not in (None, "", "None"):
            parts.append(f"resp. : {a.get('responsable')}")
        if a.get("echeance") not in (None, "", "None"):
            parts.append(f"échéance : {a.get('echeance')}")

        line = ", ".join(parts)
        try:
            doc.add_paragraph(line, style="Enumération")
        except Exception:
            doc.add_paragraph(line)

    # -------------------------------------------------
    # Perspectives
    # -------------------------------------------------
    doc.add_heading("Perspectives", level=2)
    for p in (data.get("perspectives") or []):
        if not isinstance(p, dict):
            continue
        if p.get("probleme"):
            try:
                doc.add_paragraph(p["probleme"], style="Problème")
            except Exception:
                doc.add_paragraph(p["probleme"])
        if p.get("solution"):
            try:
                doc.add_paragraph(p["solution"], style="Solution")
            except Exception:
                doc.add_paragraph(p["solution"])

    # -------------------------------------------------
    # Annexes
    # -------------------------------------------------
    doc.add_heading("Annexes", level=2)
    for ann in (data.get("annexes") or []):
        if ann is None:
            continue
        try:
            doc.add_paragraph(str(ann), style="Enumération")
        except Exception:
            doc.add_paragraph(str(ann))

    # -------------------------------------------------
    # Demandes de documents (global – dédupliquées)
    # -------------------------------------------------
    demandes = collect_global_demands(data)
    seen, uniques = set(), []

    for d in demandes:
        if not isinstance(d, dict):
            continue
        key = (
            (d.get("objet") or "").strip().lower(),
            (d.get("demandeur") or "").strip().lower(),
            (d.get("destinataire") or "").strip().lower(),
            str(d.get("echeance") or "").strip().lower(),
            str(d.get("_sujet") or "").strip().lower(),
        )
        if key in seen:
            continue
        seen.add(key)
        uniques.append(d)


    

    if uniques:
        doc.add_heading("Demandes de documents", level=2)
        for d in uniques:
            objet = (d.get("objet") or "").strip()
            if not objet:
                continue

            parts = [objet]
            for label, key in (
                ("demandeur", "demandeur"),
                ("destinataire", "destinataire"),
                ("échéance", "echeance"),
            ):
                val = d.get(key)
                if val not in (None, "", "None"):
                    parts.append(f"{label} : {val}")

            sujet_num = d.get("_sujet")
            if sujet_num not in (None, "", "None"):
                parts.append(f"sujet : {sujet_num}")

            line = parts[0] if len(parts) == 1 else f"{parts[0]} ({', '.join(parts[1:])})"
            try:
                doc.add_paragraph(line, style="Enumération")
            except Exception:
                doc.add_paragraph(line)

            commentaire = (d.get("commentaire") or "").strip()
            if commentaire:
                doc.add_paragraph(commentaire)

    # -------------------------------------------------
    # Finalisation
    # -------------------------------------------------
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

